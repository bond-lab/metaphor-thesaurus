#!/usr/bin/env python3
# /// script
# requires-python = ">=3.11"
# dependencies = ["python-docx"]
# ///
"""
Extract structured data from THE_THESAURUS.docx into thesaurus.json.

Structure detected via paragraph/run formatting:
  - Part heading:        text starts with 'Part [IVX]+'
  - Metaphor theme:      bold + centered + all-caps (consecutive lines merged)
  - Theme relationship:  centered + starts with symbol
  - Subsection heading:  underlined
  - Entry:               bold run at start = headword
"""

import json
import re
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

RELATIONSHIP_SYMBOLS = {"<", ">", "#", ">>", "^", "v", "⇔", "⟺", "||", "↔"}
WORD_CLASSES = {
    "adj", "adjphr", "adv", "advphr", "art", "cl", "excl", "idi",
    "n", "nplur", "nphr", "pr", "pref", "prphr", "pt", "v", "verg",
    "vi", "v-inf", "virec", "vtref", "vt", "prp", "pp",
}

# Regex for a single word-class token, e.g. "n", "idi(vt+adv+adv)", "(n)"
_WC_ALTS = "|".join(sorted(WORD_CLASSES, key=len, reverse=True))
_WC_TOKEN = r"(?:\(?" + r"(?:" + _WC_ALTS + r")" + r"(?:\([^)]*\))?" + r"\)?)"
# Full word class: conversion "(n)|vt" or compound "vi+adv" or simple "n"
_WC_FULL = (
    r"(?:"
    r"(?:" + _WC_TOKEN + r"\|" + _WC_TOKEN + r")"   # conversion: (n)|vt
    r"|"
    r"(?:" + _WC_TOKEN + r"(?:\+" + _WC_TOKEN + r")*)"  # simple or compound
    r")"
)
WC_AT_END = re.compile(r"^(.*?)\s*(" + _WC_FULL + r")\s*$", re.DOTALL)

PART_RE = re.compile(r"^Part\s+(?:[IVX]+|\d+)\b", re.IGNORECASE)

# Theme names that are incomplete in the source DOCX (first half of the heading missing).
# Identified by cross-referencing guide section 4 with surrounding relationships.
THEME_FIXUPS = {
    "IS COLOUR":          "RACE IS COLOUR",
    "EXPERIENCE IS":      "EXPERIENCE IS FOOD",
    "COMMUNICATION IS":   "COMMUNICATION IS FLOW",
    "UNDERSTAND/KNOW IS": "UNDERSTAND/KNOW IS SEE",
    "IS POSITION":        "JOB IS POSITION",
    "ACTIVITY/LIFE IS":   "ACTIVITY/LIFE IS PATH",
    "IS CONTAINER":       "MIND IS CONTAINER",
    "IS BUILDING":        "MIND IS BUILDING",
    "MONEY IS":           "MONEY IS FOOD",
}


def is_all_caps(text: str) -> bool:
    letters = [c for c in text if c.isalpha()]
    return bool(letters) and all(c.isupper() for c in letters)


def run_is_bold(run) -> bool:
    return bool(run.bold)


def run_is_italic(run) -> bool:
    return bool(run.italic)


def run_is_underline(run) -> bool:
    return bool(run.underline)


def para_is_centered(para) -> bool:
    if para.alignment is not None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        return para.alignment == WD_ALIGN_PARAGRAPH.CENTER
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        jc = pPr.find(qn("w:jc"))
        if jc is not None:
            return jc.get(qn("w:val")) == "center"
    return False


def para_is_underlined(para) -> bool:
    return any(run_is_underline(r) for r in para.runs if r.text.strip())


def para_full_text(para) -> str:
    return "".join(r.text for r in para.runs)


def clean_theme_name(text: str) -> str:
    """Normalise spacing artifacts and fix known incomplete theme names."""
    text = re.sub(r"/\s+", "/", text)   # 'PLACE/ LANDSCAPE' → 'PLACE/LANDSCAPE'
    text = re.sub(r"\s+/", "/", text)   # 'LANDSCAPE /BODY'  → 'LANDSCAPE/BODY'
    text = re.sub(r"\s+", " ", text)
    text = text.strip()
    return THEME_FIXUPS.get(text, text)


def starts_with_relationship_symbol(text: str) -> bool:
    t = text.strip()
    for sym in RELATIONSHIP_SYMBOLS:
        if t.startswith(sym):
            return True
    return False


def extract_relationship(text: str) -> tuple[str, str]:
    """Split '# BAD/UNIMPORTANT IS POOR/CHEAP' into ('#', 'BAD/UNIMPORTANT IS POOR/CHEAP')."""
    t = text.strip()
    for sym in sorted(RELATIONSHIP_SYMBOLS, key=len, reverse=True):
        if t.startswith(sym):
            return sym, t[len(sym):].strip()
    return "", t


def extract_relationships(text: str) -> list[tuple[str, str]]:
    """
    Parse a relationship line containing one or more comma-separated themes,
    each optionally prefixed with its own relationship symbol.

    e.g. '⇔IMPRESSIVE/FAMOUS IS LIGHT, ⇔INFORMATION IS MINERAL'
         '# BAD/UNIMPORTANT IS POOR/CHEAP, HUMAN IS VALUABLE OBJECT/COMMODITY'
    Returns list of (symbol, theme_name) tuples.
    """
    leading_sym, remainder = extract_relationship(text)

    results = []
    for chunk in remainder.split(","):
        chunk = chunk.strip().rstrip("/").strip()
        if not chunk:
            continue
        sym, theme = extract_relationship(chunk)
        if not sym:
            sym = leading_sym
        theme = clean_theme_name(theme)
        if theme:
            results.append((sym, theme))

    return results if results else [(leading_sym, clean_theme_name(remainder))]


# ---------------------------------------------------------------------------
# Domain extraction
# ---------------------------------------------------------------------------

def extract_domains(theme_name: str) -> tuple[list[str], list[str]]:
    """
    Extract target and source domain lists from a theme name.

    'QUALITY IS MONEY/WEALTH'    → (['QUALITY'], ['MONEY', 'WEALTH'])
    'BAD/UNIMPORTANT IS POOR'    → (['BAD', 'UNIMPORTANT'], ['POOR'])
    'ACTIVITY/LIFE IS PATH'      → (['ACTIVITY', 'LIFE'], ['PATH'])
    """
    parts = re.split(r"\bIS\b", theme_name, maxsplit=1)
    if len(parts) != 2:
        return [theme_name.strip()], []
    target_str, source_str = parts
    targets = [t.strip() for t in target_str.split("/") if t.strip()]
    sources = [s.strip() for s in source_str.split("/") if s.strip()]
    return targets, sources


def collect_all_domains(data: dict) -> dict:
    """
    Aggregate all unique source and target domains across every theme.
    Returns {"targets": [...], "sources": [...]} sorted alphabetically.
    """
    targets: set[str] = set()
    sources: set[str] = set()
    for part in data["parts"]:
        for theme in part["themes"]:
            t, s = extract_domains(theme["name"])
            targets.update(t)
            sources.update(s)
    return {
        "targets": sorted(targets),
        "sources": sorted(sources),
    }


# ---------------------------------------------------------------------------
# Entry parsing
# ---------------------------------------------------------------------------

def parse_entry(para) -> dict:
    """
    Parse a thesaurus entry paragraph into structured fields.

    Run formatting roles:
      bold (non-caps)  → headword
      plain            → literal meaning + word class
      ALL CAPS         → metaphorical meaning
      italic           → example sentence
    """
    entry: dict = {
        "headword": "",
        "reversal_prefix": "",
        "literal_meaning": "",
        "word_class_literal": "",
        "word_class_metaphorical": "",
        "metaphorical_meaning": "",
        "example": "",
    }

    segments = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        segments.append((run_is_bold(run), run_is_italic(run), is_all_caps(text), text))

    if not segments:
        return entry

    # Strip ">>" reversal prefix from the leading bold run
    if segments and segments[0][0]:
        first_text = segments[0][3]
        m = re.match(r"^(>>)\s*", first_text)
        if m:
            entry["reversal_prefix"] = ">>"
            segments[0] = (segments[0][0], segments[0][1], segments[0][2], first_text[m.end():])

    # Collect headword from leading bold non-caps runs
    headword_parts = []
    i = 0
    while i < len(segments):
        bold, italic, caps, text = segments[i]
        if bold and not caps:
            headword_parts.append(text)
            i += 1
        else:
            break
    entry["headword"] = "".join(headword_parts).strip()

    # Classify remaining runs
    plain_parts, caps_parts, italic_parts = [], [], []
    while i < len(segments):
        bold, italic, caps, text = segments[i]
        if italic and not caps:
            italic_parts.append(text)
        elif caps and not italic:
            caps_parts.append(text)
        elif not bold and not italic and not caps:
            plain_parts.append(text)
        else:
            (caps_parts if caps else plain_parts).append(text)
        i += 1

    plain_text = "".join(plain_parts).strip()
    entry["metaphorical_meaning"] = " ".join(caps_parts).strip()
    entry["example"] = "".join(italic_parts).strip()

    entry["literal_meaning"], entry["word_class_literal"], entry["word_class_metaphorical"] = \
        _parse_literal_and_wordclass(plain_text)

    return entry


def _parse_literal_and_wordclass(text: str) -> tuple[str, str, str]:
    """
    Split plain entry text into (literal_meaning, word_class_literal, word_class_metaphorical).

    Examples:
      "large amount of money  n"        → ("large amount of money", "n", "n")
      "(port for a ship) (n)|vt"        → ("(port for a ship)", "n", "vt")
      "(be suspended__) idi(vi+adv+adv)"→ ("(be suspended__)", "idi(vi+adv+adv)", "idi(vi+adv+adv)")
    """
    text = text.strip()
    m = WC_AT_END.match(text)
    if m:
        literal = m.group(1).strip()
        wc_str = m.group(2).strip()
    else:
        return text, "", ""

    # Parse conversion "(n)|vt" → literal_wc="n", meta_wc="vt"
    if "|" in wc_str:
        left, right = wc_str.split("|", 1)
        wc_literal = left.strip("() ")
        wc_meta = right.strip("() ")
    else:
        wc_literal = wc_meta = wc_str

    return literal, wc_literal, wc_meta


# ---------------------------------------------------------------------------
# Document classification
# ---------------------------------------------------------------------------

def classify_para(para, prev_was_theme: bool = False) -> str:
    """Return one of: 'part', 'theme', 'relationship', 'subsection', 'entry', 'blank', 'other'"""
    text = para_full_text(para).strip()
    if not text:
        return "blank"

    # Part heading: text starts with 'Part I/II/III/...' regardless of font styling
    if PART_RE.match(text):
        return "part"

    centered = para_is_centered(para)
    all_bold = all(run_is_bold(r) for r in para.runs if r.text.strip())
    all_caps_text = is_all_caps(text)

    # Theme: bold + centered + all-caps
    if centered and all_bold and all_caps_text:
        return "theme"

    # Relationship: centered + starts with a relationship symbol
    if centered and starts_with_relationship_symbol(text):
        return "relationship"

    # Centered all-caps line immediately after a theme = unlabelled relationship
    if centered and all_caps_text and prev_was_theme:
        return "relationship"

    # Subsection: underlined
    if para_is_underlined(para):
        return "subsection"

    # Entry: first non-blank run is bold
    for run in para.runs:
        if run.text.strip():
            if run_is_bold(run):
                return "entry"
            break

    return "other"


# ---------------------------------------------------------------------------
# Main parser
# ---------------------------------------------------------------------------

def parse_thesaurus(docx_path: Path) -> dict:
    doc = docx.Document(str(docx_path))

    # Pre-classify all paragraphs, merging consecutive bold+centered+all-caps
    # paragraphs into a single theme (handles theme names split across lines).
    classified: list[tuple[str, str]] = []  # (kind, text)
    prev_kind = "blank"
    for para in doc.paragraphs:
        kind = classify_para(para, prev_was_theme=(prev_kind == "theme"))
        text = para_full_text(para).strip()

        if kind == "theme" and classified and classified[-1][0] == "theme":
            # Merge with previous theme line
            prev_text = classified[-1][1]
            classified[-1] = ("theme", clean_theme_name(prev_text + " " + text))
        else:
            if kind == "theme":
                text = clean_theme_name(text)
            classified.append((kind, text))

        prev_kind = kind

    # Now build the data structure from the classified list,
    # pairing entries with their original paragraphs for run-level parsing.
    # We need the original para objects for entries, so we track index.
    para_list = list(doc.paragraphs)
    para_idx = 0

    result: dict = {"parts": [], "domains": {}}
    current_part: dict | None = None
    current_theme: dict | None = None
    current_subsection: dict | None = None

    def ensure_part():
        nonlocal current_part
        if current_part is None:
            current_part = {"name": "", "themes": []}
            result["parts"].append(current_part)

    def ensure_theme():
        nonlocal current_theme, current_subsection
        ensure_part()
        if current_theme is None:
            current_theme = {"name": "", "relationships": [], "subsections": []}
            current_part["themes"].append(current_theme)
            current_subsection = None

    def ensure_subsection(heading=""):
        nonlocal current_subsection
        ensure_theme()
        if current_subsection is None or (heading and current_subsection["heading"] != heading):
            current_subsection = {"heading": heading, "entries": []}
            current_theme["subsections"].append(current_subsection)

    # Walk classified list in sync with paragraphs
    for kind, text in classified:
        # Advance para_idx to the next non-skipped paragraph matching this kind+text
        # (needed for entry parsing which uses run data)
        while para_idx < len(para_list):
            p_text = para_full_text(para_list[para_idx]).strip()
            if p_text or kind == "blank":
                break
            para_idx += 1

        if kind == "blank":
            para_idx += 1
            continue

        elif kind == "part":
            current_part = {"name": text, "themes": []}
            result["parts"].append(current_part)
            current_theme = None
            current_subsection = None

        elif kind == "theme":
            ensure_part()
            current_theme = {"name": text, "relationships": [], "subsections": []}
            current_part["themes"].append(current_theme)
            current_subsection = None

        elif kind == "relationship":
            ensure_theme()
            for sym, theme_name in extract_relationships(text):
                current_theme["relationships"].append({"symbol": sym, "theme": theme_name})

        elif kind == "subsection":
            ensure_subsection(text)

        elif kind == "entry":
            ensure_subsection()
            # Find the matching paragraph for run-level parsing
            while para_idx < len(para_list):
                if para_full_text(para_list[para_idx]).strip() == text:
                    break
                para_idx += 1
            if para_idx < len(para_list):
                entry = parse_entry(para_list[para_idx])
                if entry["headword"] or entry["metaphorical_meaning"]:
                    current_subsection["entries"].append(entry)

        para_idx += 1

    result["domains"] = collect_all_domains(result)
    return result


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def run_tests(data: dict):
    errors = []

    def _norm(s: str) -> str:
        return " ".join(s.split()).casefold()

    def find_entry(theme_name: str, headword: str) -> dict | None:
        for part in data["parts"]:
            for theme in part["themes"]:
                if _norm(theme["name"]) == _norm(theme_name):
                    for sub in theme["subsections"]:
                        for entry in sub["entries"]:
                            if _norm(entry["headword"]) == _norm(headword):
                                return entry
        return None

    def find_theme(theme_name: str) -> dict | None:
        for part in data["parts"]:
            for theme in part["themes"]:
                if _norm(theme["name"]) == _norm(theme_name):
                    return theme
        return None

    # Test 1: basic entry
    e = find_entry("QUALITY IS MONEY/WEALTH", "wealth")
    if e is None:
        errors.append("FAIL: 'wealth' not found under QUALITY IS MONEY/WEALTH")
    elif "LARGE AMOUNT" not in e["metaphorical_meaning"]:
        errors.append("FAIL: 'wealth' metaphorical_meaning wrong: %r" % e["metaphorical_meaning"])
    else:
        print("PASS: 'wealth' entry found and parsed")

    # Test 2: example sentence
    if e and not e["example"]:
        errors.append("FAIL: 'wealth' missing example sentence")
    elif e:
        print("PASS: 'wealth' example: %r" % e["example"])

    # Test 3: word class
    if e and e["word_class_metaphorical"].strip().strip("()") != "n":
        errors.append("FAIL: 'wealth' word class wrong: %r" % e["word_class_metaphorical"])
    elif e:
        print("PASS: 'wealth' word class: %r" % e["word_class_metaphorical"])

    # Test 4: relationships split correctly
    t = find_theme("QUALITY IS MONEY/WEALTH")
    if t is None:
        errors.append("FAIL: theme QUALITY IS MONEY/WEALTH not found")
    elif len(t["relationships"]) < 2:
        errors.append("FAIL: relationships not split: %r" % t["relationships"])
    else:
        names = [r["theme"] for r in t["relationships"]]
        if not any("BAD/UNIMPORTANT" in n for n in names):
            errors.append("FAIL: expected BAD/UNIMPORTANT in relationships: %r" % names)
        elif not any("HUMAN IS VALUABLE" in n for n in names):
            errors.append("FAIL: expected HUMAN IS VALUABLE in relationships: %r" % names)
        else:
            print("PASS: relationships split: %r" % names)

    # Test 5: reversal prefix >>
    e2 = find_entry("QUALITY IS MONEY/WEALTH", "grand")
    if e2 is None:
        errors.append("FAIL: 'grand' (>> entry) not found")
    elif e2["reversal_prefix"] != ">>":
        errors.append("FAIL: 'grand' reversal_prefix wrong: %r" % e2["reversal_prefix"])
    else:
        print("PASS: 'grand' reversal_prefix=%r" % e2["reversal_prefix"])

    # Test 6: word class for simple adj entry
    e3 = find_entry("BAD/UNIMPORTANT IS POOR/CHEAP", "poor")
    if e3 is None:
        errors.append("FAIL: 'poor' not found under BAD/UNIMPORTANT IS POOR/CHEAP")
    else:
        print("PASS: 'poor' found, wc_meta=%r" % e3["word_class_metaphorical"])

    # Test 7: parts — expect 6
    n_parts = len(data["parts"])
    if n_parts < 6:
        errors.append("FAIL: expected 6 parts, got %d" % n_parts)
    else:
        print("PASS: %d parts found" % n_parts)

    # Test 8: merged theme — MIND IS CONTAINER should exist (was split across lines)
    t2 = find_theme("MIND IS CONTAINER")
    if t2 is None:
        errors.append("FAIL: 'MIND IS CONTAINER' theme not found (merge failed?)")
    else:
        print("PASS: 'MIND IS CONTAINER' theme found (consecutive lines merged)")

    # Test 9: idi word class parses fully, e.g. hang in there
    e4 = find_entry("AVOID FAILURE", "hang in there")
    if e4 is None:
        # search more broadly
        for part in data["parts"]:
            for theme in part["themes"]:
                for sub in theme["subsections"]:
                    for entry in sub["entries"]:
                        if entry["headword"].startswith("hang in there"):
                            e4 = entry
                            break
    if e4 is None:
        errors.append("FAIL: 'hang in there' not found")
    elif not e4["word_class_metaphorical"].startswith("idi"):
        errors.append("FAIL: 'hang in there' word class wrong: %r" % e4["word_class_metaphorical"])
    elif "(" not in e4["word_class_metaphorical"]:
        errors.append("FAIL: 'hang in there' word class truncated (missing parens): %r" % e4["word_class_metaphorical"])
    else:
        print("PASS: 'hang in there' word class: %r" % e4["word_class_metaphorical"])

    # Test 10: domains extracted
    domains = data.get("domains", {})
    if not domains.get("targets") or not domains.get("sources"):
        errors.append("FAIL: domains not extracted")
    else:
        print("PASS: %d target domains, %d source domains" % (
            len(domains["targets"]), len(domains["sources"])))

    # Test 11: total entry count
    total = sum(
        len(sub["entries"])
        for part in data["parts"]
        for theme in part["themes"]
        for sub in theme["subsections"]
    )
    if total < 100:
        errors.append("FAIL: only %d entries found — expected thousands" % total)
    else:
        print("PASS: %d total entries" % total)

    return errors


# ---------------------------------------------------------------------------
# Summary
# ---------------------------------------------------------------------------

def summarize(data: dict):
    from collections import Counter

    all_themes = [
        (part["name"], theme)
        for part in data["parts"]
        for theme in part["themes"]
    ]
    all_entries = [
        entry
        for _, theme in all_themes
        for sub in theme["subsections"]
        for entry in sub["entries"]
    ]
    all_relationships = [
        rel
        for _, theme in all_themes
        for rel in theme["relationships"]
    ]

    total_entries = len(all_entries)
    total_themes = len(all_themes)
    total_subsections = sum(len(t["subsections"]) for _, t in all_themes)
    total_relationships = len(all_relationships)

    print("=" * 60)
    print("THESAURUS SUMMARY")
    print("=" * 60)
    print("  Parts:         %d" % len(data["parts"]))
    print("  Themes:        %d" % total_themes)
    print("  Subsections:   %d" % total_subsections)
    print("  Entries:       %d" % total_entries)
    print("  Relationships: %d" % total_relationships)

    domains = data.get("domains", {})
    print("  Target domains: %d" % len(domains.get("targets", [])))
    print("  Source domains: %d" % len(domains.get("sources", [])))

    print("\n--- Entries and themes by part ---")
    for part in data["parts"]:
        n_themes = len(part["themes"])
        n_entries = sum(
            len(sub["entries"])
            for theme in part["themes"]
            for sub in theme["subsections"]
        )
        name = part["name"] or "(unnamed)"
        print("  %s" % name)
        print("    %d themes, %d entries" % (n_themes, n_entries))

    print("\n--- Top 10 themes by entry count ---")
    theme_counts = sorted(
        [(t["name"], sum(len(s["entries"]) for s in t["subsections"])) for _, t in all_themes],
        key=lambda x: x[1], reverse=True
    )
    for name, count in theme_counts[:10]:
        print("  %4d  %s" % (count, name))

    print("\n--- Top 10 themes by relationship count ---")
    rel_counts = sorted(
        [(t["name"], len(t["relationships"])) for _, t in all_themes],
        key=lambda x: x[1], reverse=True
    )
    for name, count in rel_counts[:10]:
        if count > 0:
            print("  %3d  %s" % (count, name))

    print("\n--- Word class distribution (metaphorical, top 15) ---")
    wc_counter: Counter = Counter()
    for e in all_entries:
        wc = e.get("word_class_metaphorical", "").strip().strip("()")
        if wc:
            wc_counter[wc] += 1
    for wc, count in wc_counter.most_common(15):
        pct = 100 * count / total_entries
        print("  %-14s  %5d  (%.1f%%)" % (wc, count, pct))

    print("\n--- Relationship symbol distribution ---")
    sym_counter: Counter = Counter(r["symbol"] or "(none)" for r in all_relationships)
    for sym, count in sym_counter.most_common():
        print("  %-8s  %d" % (sym, count))

    print("\n--- Top 20 target domains ---")
    target_counter: Counter = Counter()
    source_counter: Counter = Counter()
    for _, theme in all_themes:
        ts, ss = extract_domains(theme["name"])
        target_counter.update(ts)
        source_counter.update(ss)
    for domain, count in target_counter.most_common(20):
        print("  %4d  %s" % (count, domain))

    print("\n--- Top 20 source domains ---")
    for domain, count in source_counter.most_common(20):
        print("  %4d  %s" % (count, domain))

    reversals = [e for e in all_entries if e.get("reversal_prefix") == ">>"]
    no_example = [e for e in all_entries if not e.get("example")]
    print("\n  Entries with '>>' prefix:         %d" % len(reversals))
    print("  Entries without example sentence: %d" % len(no_example))
    print("=" * 60)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    docx_path = Path("THE_THESAURUS.docx")
    if not docx_path.exists():
        print("Error: %s not found" % docx_path, file=sys.stderr)
        sys.exit(1)

    print("Parsing %s ..." % docx_path)
    data = parse_thesaurus(docx_path)

    print("\nRunning tests ...")
    errors = run_tests(data)

    if errors:
        print("\nTest failures:")
        for e in errors:
            print("  " + e)
        sys.exit(1)

    out_path = Path("thesaurus.json")
    with out_path.open("w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print("\nWrote %s (%d KB)" % (out_path, out_path.stat().st_size // 1024))

    print()
    summarize(data)
    print("\nAll tests passed.")


if __name__ == "__main__":
    main()
